# # importing required classes
# from pypdf import PdfReader

# # creating a pdf reader object
# reader = PdfReader(r'C:\Users\aravi\OneDrive\Documents\Resume_Scanner_Bot\data\New-York-Resume-Template-Creative.pdf')

# # printing number of pages in pdf file
# print(len(reader.pages))

# # creating a page object
# page = reader.pages[0]

# # extracting text from page
# print(page.extract_text())




# # import the Document class 
# # from the docx module
# from docx import Document

# # create an instance of a 
# # word document we want to open
# doc = Document(r'C:\Users\aravi\OneDrive\Documents\Resume_Scanner_Bot\data\JD_sd.docx')

# # print the list of paragraphs in the document
# print('List of paragraph objects:->>>')
# print(doc.paragraphs)

# # print the list of the runs 
# # in a specified paragraph
# print('\nList of runs objects in 1st paragraph:->>>')
# print(doc.paragraphs[0].runs)

# # print the text in a paragraph 
# print('\nText in the 1st paragraph:->>>')
# print(doc.paragraphs[0].text)

# # for printing the complete document
# print('\nThe whole content of the document:->>>\n')
# for para in doc.paragraphs:
#     print(para.text)













#####     Here this code part is to read the resume and jd in api not checking the percentage and giving report#####

# from fastapi import FastAPI, UploadFile, File
# from fastapi.responses import JSONResponse
# from pypdf import PdfReader
# from docx import Document
# import io

# app = FastAPI(title="Resume and JD Extractor API")

# def extract_text_from_pdf(file_bytes):
#     """Extract text from PDF bytes."""
#     reader = PdfReader(io.BytesIO(file_bytes))
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text() or ""
#     return text.strip()

# def extract_text_from_docx(file_bytes):
#     """Extract text from DOCX bytes."""
#     doc = Document(io.BytesIO(file_bytes))
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text.strip()

# @app.post("/extract/")
# async def extract_text(
#     resume: UploadFile = File(...),
#     jd: UploadFile = File(...)
# ):
#     """Upload resume and job description; extract text and return as JSON."""

#     try:
#         # --- Extract Resume Text ---
#         resume_bytes = await resume.read()
#         if resume.filename.endswith(".pdf"):
#             resume_text = extract_text_from_pdf(resume_bytes)
#         elif resume.filename.endswith(".docx"):
#             resume_text = extract_text_from_docx(resume_bytes)
#         else:
#             return JSONResponse({"error": "Resume must be PDF or DOCX"}, status_code=400)

#         # --- Extract JD Text ---
#         jd_bytes = await jd.read()
#         if jd.filename.endswith(".pdf"):
#             jd_text = extract_text_from_pdf(jd_bytes)
#         elif jd.filename.endswith(".docx"):
#             jd_text = extract_text_from_docx(jd_bytes)
#         else:
#             return JSONResponse({"error": "JD must be PDF or DOCX"}, status_code=400)

#         # --- Return Extracted Data ---
#         return {
#             "status": "success",
#             "resume_filename": resume.filename,
#             "jd_filename": jd.filename,
#             "resume_text": resume_text[:1000],  # limit preview
#             "jd_text": jd_text[:1000]
#         }

#     except Exception as e:
#         return JSONResponse({"error": str(e)}, status_code=500)






########### now executing my code#########################  


# from fastapi import FastAPI, UploadFile, File
# from fastapi.responses import JSONResponse
# from pypdf import PdfReader
# from docx import Document
# import io
# import json
# from groq import Groq

# # ------------------- CONFIGURE GROQ -------------------
# # Option 1: Directly in code (for testing)
# groq_api_key = "
# client = Groq(api_key=groq_api_key)


#     for page in reader.pages:
#         text += page.extract_text() or ""
#     return text.strip()

# def extract_text_from_docx(file_bytes):
#     doc = Document(io.BytesIO(file_bytes))
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text.strip()

# # ------------------- AI ANALYSIS -------------------
# def analyze_with_groq(resume_text, jd_text):
#     prompt = f"""
# You are an expert recruiter.

# Compare the following Job Description and Resume.

# Job Description:
# {jd_text}

# Resume:
# {resume_text}

# Tasks:
# 1. Provide a match percentage (0-100).
# 2. List key skills matched.
# 3. List missing or weak skills.
# 4. Explain why the score is high or low.

# IMPORTANT: Return output as **valid JSON ONLY**, strictly in this format:

# {{
#   "match_score": <integer>,
#   "matched_skills": ["skill1", "skill2"],
#   "missing_skills": ["skill1", "skill2"],
#   "summary": "<explanation>"
# }}
# """
#     response = client.chat.completions.create(
#         messages=[{"role": "user", "content": prompt}],
#         model="llama-3.3-70b-versatile",
#     )

#     ai_output = response.choices[0].message.content.strip()

#     # --- Remove backticks or markdown code blocks if present ---
#     if ai_output.startswith("```"):
#         ai_output = ai_output.strip("`")  # removes leading/trailing backticks
#         # Sometimes the AI puts ```json or ``` at the start, remove that too
#         if ai_output.lower().startswith("json"):
#             ai_output = ai_output[4:].strip()

#     # --- Parse JSON safely ---
#     try:
#         result = json.loads(ai_output)
#     except json.JSONDecodeError:
#         # Return raw output if JSON parsing fails
#         result = {
#             "match_score": 0,
#             "matched_skills": [],
#             "missing_skills": [],
#             "summary": f"⚠️ AI did not return valid JSON. Raw output: {ai_output}"
#         }
#     return result


# # ------------------- API ENDPOINT -------------------
# @app.post("/analyze/")
# async def analyze_resume_and_jd(resume: UploadFile = File(...), jd: UploadFile = File(...)):
#     try:
#         # --- Extract Resume Text ---
#         resume_bytes = await resume.read()
#         if resume.filename.endswith(".pdf"):
#             resume_text = extract_text_from_pdf(resume_bytes)
#         elif resume.filename.endswith(".docx"):
#             resume_text = extract_text_from_docx(resume_bytes)
#         else:
#             return JSONResponse({"error": "Resume must be PDF or DOCX"}, status_code=400)

#         # --- Extract JD Text ---
#         jd_bytes = await jd.read()
#         if jd.filename.endswith(".pdf"):
#             jd_text = extract_text_from_pdf(jd_bytes)
#         elif jd.filename.endswith(".docx"):
#             jd_text = extract_text_from_docx(jd_bytes)
#         else:
#             return JSONResponse({"error": "JD must be PDF or DOCX"}, status_code=400)

#         # --- Analyze using Groq AI ---
#         ai_result = analyze_with_groq(resume_text, jd_text)

#         # --- 80% Threshold Check ---
#         if ai_result["match_score"] >= 80:
#             ai_result["summary"] = "✅ Strong match: " + ai_result["summary"]
#         else:
#             ai_result["summary"] = "⚠️ Weak match: " + ai_result["summary"]

#         # --- Return Report ---
#         report = {
#             "resume_filename": resume.filename,
#             "jd_filename": jd.filename,
#             "match_score": ai_result["match_score"],
#             "matched_skills": ai_result["matched_skills"],
#             "missing_skills": ai_result["missing_skills"],
#             "summary": ai_result["summary"]
#         }

#         return report

#     except Exception as e:
#         return JSONResponse({"error": str(e)}, status_code=500)







################       hideen api key code #################



# import os
# from fastapi import FastAPI, UploadFile, File
# from fastapi.responses import JSONResponse
# from pypdf import PdfReader
# from docx import Document
# import io
# import json
# from groq import Groq
# import os
# from dotenv import load_dotenv

# # ------------------- LOAD ENVIRONMENT VARIABLES -------------------
# load_dotenv()  # Reads the .env file in the project root

# groq_api_key = os.getenv("GROQ_API_KEY")  # Load your API key from .env
# if not groq_api_key:
#     raise ValueError("GROQ_API_KEY not found. Please set it in .env")

# client = Groq(api_key=groq_api_key)

# # ------------------- INITIALIZE FASTAPI -------------------
# app = FastAPI(title="Resume and JD Analyzer API")

# # ------------------- TEXT EXTRACTION -------------------
# def extract_text_from_pdf(file_bytes):
#     reader = PdfReader(io.BytesIO(file_bytes))
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text() or ""
#     return text.strip()

# def extract_text_from_docx(file_bytes):
#     doc = Document(io.BytesIO(file_bytes))
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text.strip()

# # ------------------- AI ANALYSIS -------------------
# def analyze_with_groq(resume_text, jd_text):
#     prompt = f"""
# You are an expert recruiter.

# Compare the following Job Description and Resume.

# Job Description:
# {jd_text}

# Resume:
# {resume_text}

# Tasks:
# 1. Provide a match percentage (0-100).
# 2. List key skills matched.
# 3. List missing or weak skills.
# 4. Explain why the score is high or low.

# IMPORTANT: Return output as **valid JSON ONLY**, strictly in this format:

# {{
#   "match_score": <integer>,
#   "matched_skills": ["skill1", "skill2"],
#   "missing_skills": ["skill1", "skill2"],
#   "summary": "<explanation>"
# }}
# """
#     response = client.chat.completions.create(
#         messages=[{"role": "user", "content": prompt}],
#         model="llama-3.3-70b-versatile",
#     )

#     ai_output = response.choices[0].message.content.strip()

#     if ai_output.startswith("```"):
#         ai_output = ai_output.strip("`")
#         if ai_output.lower().startswith("json"):
#             ai_output = ai_output[4:].strip()

#     try:
#         result = json.loads(ai_output)
#     except json.JSONDecodeError:
#         result = {
#             "match_score": 0,
#             "matched_skills": [],
#             "missing_skills": [],
#             "summary": f"⚠️ AI did not return valid JSON. Raw output: {ai_output}"
#         }
#     return result

# # ------------------- API ENDPOINT -------------------
# @app.post("/analyze/")
# async def analyze_resume_and_jd(resume: UploadFile = File(...), jd: UploadFile = File(...)):
#     try:
#         resume_bytes = await resume.read()
#         if resume.filename.endswith(".pdf"):
#             resume_text = extract_text_from_pdf(resume_bytes)
#         elif resume.filename.endswith(".docx"):
#             resume_text = extract_text_from_docx(resume_bytes)
#         else:
#             return JSONResponse({"error": "Resume must be PDF or DOCX"}, status_code=400)

#         jd_bytes = await jd.read()
#         if jd.filename.endswith(".pdf"):
#             jd_text = extract_text_from_pdf(jd_bytes)
#         elif jd.filename.endswith(".docx"):
#             jd_text = extract_text_from_docx(jd_bytes)
#         else:
#             return JSONResponse({"error": "JD must be PDF or DOCX"}, status_code=400)

#         ai_result = analyze_with_groq(resume_text, jd_text)

#         if ai_result["match_score"] >= 80:
#             ai_result["summary"] = "✅ Strong match: " + ai_result["summary"]
#         else:
#             ai_result["summary"] = "⚠️ Weak match: " + ai_result["summary"]

#         report = {
#             "resume_filename": resume.filename,
#             "jd_filename": jd.filename,
#             "match_score": ai_result["match_score"],
#             "matched_skills": ai_result["matched_skills"],
#             "missing_skills": ai_result["missing_skills"],
#             "summary": ai_result["summary"]
#         }

#         return report

#     except Exception as e:
#         return JSONResponse({"error": str(e)}, status_code=500)





##############################   the updated code for creating link with render web app###################



import os
import io
import json
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import JSONResponse
from pypdf import PdfReader
from docx import Document
from groq import Groq
from dotenv import load_dotenv

# ------------------- LOAD ENVIRONMENT VARIABLES -------------------
load_dotenv()  # Reads .env file if present

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("❌ GROQ_API_KEY not found. Please add it to your Render environment or .env file.")

client = Groq(api_key=GROQ_API_KEY)

# ------------------- INITIALIZE FASTAPI -------------------
app = FastAPI(title="Resume and JD Analyzer API")

# ------------------- TEXT EXTRACTION -------------------
def extract_text_from_pdf(file_bytes):
    """Extract text from PDF files."""
    reader = PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text.strip()


def extract_text_from_docx(file_bytes):
    """Extract text from DOCX files."""
    doc = Document(io.BytesIO(file_bytes))
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

# ------------------- AI ANALYSIS -------------------
def analyze_with_groq(resume_text, jd_text):
    """Send texts to Groq model and get analysis."""
    prompt = f"""
You are an expert recruiter.

Compare the following Job Description and Resume.

Job Description:
{jd_text}

Resume:
{resume_text}

Tasks:
1. Provide a match percentage (0-100).
2. List key skills matched.
3. List missing or weak skills.
4. Explain why the score is high or low.

Return output as **valid JSON ONLY** in this format:

{{
  "match_score": <integer>,
  "matched_skills": ["skill1", "skill2"],
  "missing_skills": ["skill1", "skill2"],
  "summary": "<explanation>"
}}
"""

    response = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )

    ai_output = response.choices[0].message.content.strip()

    # Clean output if wrapped in code block
    if ai_output.startswith("```"):
        ai_output = ai_output.strip("`")
        if ai_output.lower().startswith("json"):
            ai_output = ai_output[4:].strip()

    try:
        result = json.loads(ai_output)
    except json.JSONDecodeError:
        result = {
            "match_score": 0,
            "matched_skills": [],
            "missing_skills": [],
            "summary": f"⚠️ AI returned invalid JSON. Raw output: {ai_output}",
        }

    return result

# ------------------- API ENDPOINT -------------------
@app.post("/analyze/")
async def analyze_resume_and_jd(resume: UploadFile = File(...), jd: UploadFile = File(...)):
    """Main endpoint to analyze resume and job description."""
    try:
        # --- Extract Resume ---
        resume_bytes = await resume.read()
        if resume.filename.endswith(".pdf"):
            resume_text = extract_text_from_pdf(resume_bytes)
        elif resume.filename.endswith(".docx"):
            resume_text = extract_text_from_docx(resume_bytes)
        else:
            return JSONResponse({"error": "Resume must be PDF or DOCX"}, status_code=400)

        # --- Extract Job Description ---
        jd_bytes = await jd.read()
        if jd.filename.endswith(".pdf"):
            jd_text = extract_text_from_pdf(jd_bytes)
        elif jd.filename.endswith(".docx"):
            jd_text = extract_text_from_docx(jd_bytes)
        else:
            return JSONResponse({"error": "JD must be PDF or DOCX"}, status_code=400)

        # --- Analyze with Groq ---
        ai_result = analyze_with_groq(resume_text, jd_text)

        # --- Improve Summary ---
        if ai_result["match_score"] >= 80:
            ai_result["summary"] = "✅ Strong match: " + ai_result["summary"]
        else:
            ai_result["summary"] = "⚠️ Weak match: " + ai_result["summary"]

        report = {
            "resume_filename": resume.filename,
            "jd_filename": jd.filename,
            "match_score": ai_result["match_score"],
            "matched_skills": ai_result["matched_skills"],
            "missing_skills": ai_result["missing_skills"],
            "summary": ai_result["summary"],
        }

        return JSONResponse(content=report)

    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

# ------------------- HEALTH CHECK (for Render) -------------------
@app.get("/")
def root():
    return {"message": "✅ Resume Analyzer API is running successfully!"}

# ------------------- RUN LOCALLY -------------------
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8000)))
